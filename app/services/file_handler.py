"""
File Handler Service - Document processing and text extraction
Supports multiple file formats with OCR capabilities
"""
import os
import uuid
import mimetypes
from pathlib import Path
from typing import Tuple, Optional, Dict, Any
import structlog

import PyPDF2
from PIL import Image
import pytesseract
from docx import Document

from config import get_settings

logger = structlog.get_logger(__name__)


class FileHandler:
    """
    Handle file upload, validation, and text extraction
    Supports PDF, DOCX, TXT, RTF, and image files with OCR
    """
    
    def __init__(self):
        self.settings = get_settings()
        self.upload_folder = Path(self.settings.UPLOAD_FOLDER)
        self.upload_folder.mkdir(parents=True, exist_ok=True)
        
        # Supported file types
        self.supported_extensions = {
            '.pdf': self._extract_from_pdf,
            '.txt': self._extract_from_text,
            '.rtf': self._extract_from_text,
            '.docx': self._extract_from_docx,
            '.doc': self._extract_from_docx,
            '.png': self._extract_from_image,
            '.jpg': self._extract_from_image,
            '.jpeg': self._extract_from_image,
            '.tiff': self._extract_from_image,
            '.bmp': self._extract_from_image
        }
    
    def validate_file(self, file_data: bytes, filename: str) -> Tuple[bool, str]:
        """
        Validate uploaded file
        
        Args:
            file_data: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            # Check file size
            if len(file_data) > self.settings.max_file_size_bytes:
                return False, f"File size {len(file_data):,} bytes exceeds limit of {self.settings.max_file_size_bytes:,} bytes"
            
            # Check file extension
            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.supported_extensions:
                supported = ", ".join(self.supported_extensions.keys())
                return False, f"File type '{file_ext}' not supported. Supported types: {supported}"
            
            # Check if file is empty
            if len(file_data) == 0:
                return False, "File is empty"
            
            # Validate file content matches extension
            if not self._validate_file_content(file_data, file_ext):
                return False, f"File content doesn't match extension {file_ext}"
            
            return True, ""
            
        except Exception as e:
            logger.error("File validation failed", filename=filename, error=str(e))
            return False, f"Validation error: {str(e)}"
    
    def save_file(self, file_data: bytes, original_filename: str) -> Tuple[Optional[str], Optional[str]]:
        """
        Save uploaded file to disk
        
        Args:
            file_data: File content as bytes
            original_filename: Original filename from upload
            
        Returns:
            Tuple of (saved_file_path, error_message)
        """
        try:
            # Generate unique filename
            file_ext = Path(original_filename).suffix.lower()
            unique_filename = f"{uuid.uuid4()}{file_ext}"
            file_path = self.upload_folder / unique_filename
            
            # Save file
            with open(file_path, 'wb') as f:
                f.write(file_data)
            
            logger.info("File saved successfully", 
                       original_filename=original_filename,
                       saved_path=str(file_path),
                       size_bytes=len(file_data))
            
            return str(file_path), None
            
        except Exception as e:
            logger.error("File save failed", 
                        filename=original_filename, 
                        error=str(e))
            return None, f"Failed to save file: {str(e)}"
    
    def extract_text(self, file_path: str) -> Tuple[Optional[str], Optional[str], Dict[str, Any]]:
        """
        Extract text content from file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (extracted_text, error_message, metadata)
        """
        try:
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                return None, f"File not found: {file_path}", {}
            
            file_ext = file_path_obj.suffix.lower()
            
            if file_ext not in self.supported_extensions:
                return None, f"Unsupported file type: {file_ext}", {}
            
            # Extract text using appropriate method
            extractor_func = self.supported_extensions[file_ext]
            text_content, metadata = extractor_func(file_path)
            
            if not text_content or len(text_content.strip()) == 0:
                return None, "No text content could be extracted from file", metadata
            
            # Clean and normalize text
            cleaned_text = self._clean_extracted_text(text_content)
            
            logger.info("Text extraction successful",
                       file_path=file_path,
                       text_length=len(cleaned_text),
                       method=extractor_func.__name__)
            
            return cleaned_text, None, metadata
            
        except Exception as e:
            logger.error("Text extraction failed", 
                        file_path=file_path, 
                        error=str(e))
            return None, f"Text extraction error: {str(e)}", {}
    
    def _extract_from_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file"""
        text = ""
        metadata = {"pages": 0, "extraction_method": "PyPDF2"}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata["title"] = pdf_reader.metadata.get("/Title", "")
                    metadata["author"] = pdf_reader.metadata.get("/Author", "")
                    metadata["subject"] = pdf_reader.metadata.get("/Subject", "")
            
            # If no text extracted, try OCR
            if len(text.strip()) < 100:
                logger.info("PDF text extraction yielded minimal content, trying OCR")
                ocr_text, ocr_metadata = self._extract_pdf_with_ocr(file_path)
                if ocr_text and len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    metadata.update(ocr_metadata)
                    metadata["extraction_method"] = "OCR"
            
        except Exception as e:
            logger.warning("PyPDF2 extraction failed, trying OCR", error=str(e))
            # Fallback to OCR
            text, metadata = self._extract_pdf_with_ocr(file_path)
            metadata["extraction_method"] = "OCR_fallback"
        
        return text.strip(), metadata
    
    def _extract_from_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            text = "\n".join(text_parts)
            
            # Extract metadata
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "extraction_method": "python-docx"
            }
            
            # Core properties
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata["title"] = props.title or ""
                metadata["author"] = props.author or ""
                metadata["subject"] = props.subject or ""
                metadata["created"] = str(props.created) if props.created else ""
                metadata["modified"] = str(props.modified) if props.modified else ""
            
            return text, metadata
            
        except Exception as e:
            logger.error("DOCX extraction failed", file_path=file_path, error=str(e))
            return "", {"error": str(e), "extraction_method": "failed"}
    
    def _extract_from_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from text files (TXT, RTF)"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                
                metadata = {
                    "encoding": encoding,
                    "extraction_method": "text_file",
                    "lines": len(text.split('\n'))
                }
                
                return text, metadata
                
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error("Text file extraction failed", 
                           file_path=file_path, 
                           encoding=encoding, 
                           error=str(e))
        
        return "", {"error": "Could not decode file with any supported encoding"}
    
    def _extract_from_image(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from image using OCR"""
        try:
            # Open and process image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Perform OCR
            text = pytesseract.image_to_string(image, config='--psm 6')
            
            # Get additional info
            image_info = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            
            metadata = {
                "extraction_method": "tesseract_ocr",
                "image_size": image.size,
                "image_mode": image.mode,
                "confidence_scores": [int(conf) for conf in image_info['conf'] if int(conf) > 0],
                "detected_words": len([word for word in image_info['text'] if word.strip()])
            }
            
            # Calculate average confidence
            if metadata["confidence_scores"]:
                metadata["avg_confidence"] = sum(metadata["confidence_scores"]) / len(metadata["confidence_scores"])
            
            return text, metadata
            
        except Exception as e:
            logger.error("OCR extraction failed", file_path=file_path, error=str(e))
            return "", {"error": str(e), "extraction_method": "ocr_failed"}
    
    def _extract_pdf_with_ocr(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using OCR (fallback method)"""
        try:
            # This is a simplified OCR approach for PDFs
            # In production, you'd want to convert PDF pages to images first
            # then apply OCR to each page
            
            # For now, return empty with appropriate metadata
            metadata = {
                "extraction_method": "pdf_ocr",
                "note": "PDF OCR not fully implemented - would need pdf2image library",
                "pages": 0
            }
            
            return "", metadata
            
        except Exception as e:
            logger.error("PDF OCR extraction failed", file_path=file_path, error=str(e))
            return "", {"error": str(e), "extraction_method": "pdf_ocr_failed"}
    
    def _validate_file_content(self, file_data: bytes, expected_ext: str) -> bool:
        """Validate file content matches expected extension"""
        try:
            # Simple magic number checking
            if expected_ext == '.pdf':
                return file_data.startswith(b'%PDF-')
            elif expected_ext in ['.jpg', '.jpeg']:
                return file_data.startswith(b'\xff\xd8\xff')
            elif expected_ext == '.png':
                return file_data.startswith(b'\x89PNG\r\n\x1a\n')
            elif expected_ext == '.docx':
                return b'PK\x03\x04' in file_data[:4]  # ZIP signature (DOCX is a ZIP)
            elif expected_ext in ['.txt', '.rtf']:
                # Try to decode as text
                try:
                    file_data.decode('utf-8')
                    return True
                except UnicodeDecodeError:
                    try:
                        file_data.decode('latin-1')
                        return True
                    except UnicodeDecodeError:
                        return False
            
            # Default: assume valid if we can't check
            return True
            
        except Exception:
            return True  # Assume valid if validation fails
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = []
        for line in text.split('\n'):
            cleaned_line = ' '.join(line.split())
            if cleaned_line:  # Skip empty lines
                lines.append(cleaned_line)
        
        # Join with single newlines and limit length
        cleaned_text = '\n'.join(lines)
        
        # Truncate if too long (to prevent token limit issues)
        max_length = 50000  # ~50KB of text
        if len(cleaned_text) > max_length:
            cleaned_text = cleaned_text[:max_length] + "\n\n[... content truncated due to length ...]"
        
        return cleaned_text
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get comprehensive file information"""
        try:
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                return {"error": "File not found"}
            
            stat = file_path_obj.stat()
            
            return {
                "filename": file_path_obj.name,
                "extension": file_path_obj.suffix.lower(),
                "size_bytes": stat.st_size,
                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                "created_at": stat.st_ctime,
                "modified_at": stat.st_mtime,
                "mime_type": mimetypes.guess_type(str(file_path_obj))[0],
                "is_supported": file_path_obj.suffix.lower() in self.supported_extensions
            }
            
        except Exception as e:
            return {"error": str(e)}
    
    def cleanup_old_files(self, days_old: int = 30) -> int:
        """Clean up old uploaded files"""
        import time
        
        cleaned_count = 0
        cutoff_time = time.time() - (days_old * 24 * 60 * 60)
        
        try:
            for file_path in self.upload_folder.iterdir():
                if file_path.is_file() and file_path.stat().st_mtime < cutoff_time:
                    try:
                        file_path.unlink()
                        cleaned_count += 1
                        logger.info("Cleaned up old file", file_path=str(file_path))
                    except Exception as e:
                        logger.warning("Failed to clean up file", 
                                     file_path=str(file_path), 
                                     error=str(e))
            
            logger.info(f"Cleanup completed", files_cleaned=cleaned_count)
            return cleaned_count
            
        except Exception as e:
            logger.error("Cleanup operation failed", error=str(e))
            return 0